"""
Generate Technical Document: PLCVS Audio Processing Pipeline
Noise Reduction, Speech/Non-Speech Separation, Word Segmentation
- No speaker/turn detection content
- Separate diagram for each section
- Clear, humanized tone
"""

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch
import os

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUT_DIR = "/home/charlie/Desktop/DRDO"

# Colors
NAVY   = '#1B3A5C'
BLUE   = '#2E86AB'
GREEN  = '#28A745'
ORANGE = '#E8850C'
LIGHT  = '#E8F4F8'
WHITE  = '#FFFFFF'
DARK   = '#333333'
RED    = '#DC3545'


def shade_cell(cell, hex_color):
    tc = cell._element.get_or_add_tcPr()
    el = tc.makeelement(qn('w:shd'), {qn('w:fill'): hex_color.replace('#',''), qn('w:val'): 'clear'})
    tc.append(el)


def para(doc, text, size=10, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, after=6, before=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = 'Calibri'
    if bold: r.bold = True
    if color: r.font.color.rgb = RGBColor.from_string(color.replace('#',''))
    return p


def bullet(doc, text, size=9.5, after=3):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.left_indent = Cm(1.5)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = 'Calibri'
    return p


def add_image_centered(doc, path, width=Inches(6.2)):
    doc.add_picture(path, width=width)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def rounded_box(ax, x, y, w, h, text_lines, bg, edge, lw=1.8, title_size=9, body_size=7.5, title_color=None):
    """Draw a rounded box with title (first line bold) and body lines."""
    box = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.12",
                          facecolor=bg, edgecolor=edge, linewidth=lw)
    ax.add_patch(box)
    cx = x + w / 2
    if not title_color:
        title_color = edge
    # Title is first line
    top = y + h - 0.3
    ax.text(cx, top, text_lines[0], fontsize=title_size, fontweight='bold',
            ha='center', va='center', color=title_color)
    for i, line in enumerate(text_lines[1:], 1):
        ax.text(cx, top - i * 0.28, line, fontsize=body_size,
                ha='center', va='center', color=DARK)


def arrow_down(ax, x, y_from, y_to, color=NAVY, lw=1.5):
    ax.annotate('', xy=(x, y_to), xytext=(x, y_from),
                arrowprops=dict(arrowstyle='->', color=color, lw=lw))


def arrow_right(ax, x_from, y, x_to, color=NAVY, lw=1.5):
    ax.annotate('', xy=(x_to, y), xytext=(x_from, y),
                arrowprops=dict(arrowstyle='->', color=color, lw=lw))


# ═══════════════════════════════════════════════════════════
# DIAGRAM 1: High-Level Pipeline Overview
# ═══════════════════════════════════════════════════════════
def diagram_overview():
    fig, ax = plt.subplots(figsize=(9, 4.2))
    ax.set_xlim(0, 9); ax.set_ylim(0, 4.2)
    ax.axis('off'); fig.patch.set_facecolor('white')

    ax.text(4.5, 3.95, 'PLCVS: End-to-End Audio Processing Pipeline', fontsize=13,
            fontweight='bold', ha='center', color=NAVY)

    # 4 big boxes left-to-right
    specs = [
        (0.15, 'Raw Audio\nCapture',      'PyAudio\n16kHz, 16-bit PCM',    LIGHT,    NAVY),
        (2.35, 'Noise\nReduction',         'Dual-Model VAD\n(Neural + Energy)',  '#E8F8E8', GREEN),
        (4.55, 'Speech\nSegmentation',     'Stateful detector\nClean boundaries', '#E8F8E8', GREEN),
        (6.75, 'Word Segmentation\n& Post-Processing', 'Domain corrections\nTimestamp alignment', '#E8F8E8', GREEN),
    ]
    for i, (x, title, body, bg, edge) in enumerate(specs):
        rounded_box(ax, x, 1.2, 2.0, 2.2, [title, '', body], bg, edge, title_size=10, body_size=8)
        # step number circle
        cx = x + 1.0
        circle = plt.Circle((cx, 3.55), 0.18, facecolor=edge, edgecolor='white', lw=1.2)
        ax.add_patch(circle)
        ax.text(cx, 3.55, str(i+1), fontsize=10, fontweight='bold', ha='center', va='center', color='white')
        # arrow to next
        if i < 3:
            arrow_right(ax, x + 2.0, 2.3, x + 2.2, color=NAVY, lw=1.8)

    # Legend
    ax.add_patch(FancyBboxPatch((0.3, 0.3), 0.5, 0.3, boxstyle="round,pad=0.05",
                                 facecolor='#E8F8E8', edgecolor=GREEN, lw=1.5))
    ax.text(1.0, 0.45, '= Our custom contribution', fontsize=8, ha='left', va='center', color=GREEN)
    ax.add_patch(FancyBboxPatch((4.0, 0.3), 0.5, 0.3, boxstyle="round,pad=0.05",
                                 facecolor=LIGHT, edgecolor=NAVY, lw=1.5))
    ax.text(4.7, 0.45, '= Standard infrastructure', fontsize=8, ha='left', va='center', color=NAVY)

    plt.tight_layout()
    p = os.path.join(OUT_DIR, '_d1.png')
    plt.savefig(p, dpi=220, bbox_inches='tight', facecolor='white'); plt.close()
    return p


# ═══════════════════════════════════════════════════════════
# DIAGRAM 2: Noise Reduction Detail
# ═══════════════════════════════════════════════════════════
def diagram_noise_reduction():
    fig, ax = plt.subplots(figsize=(9, 5.5))
    ax.set_xlim(0, 9); ax.set_ylim(0, 5.5)
    ax.axis('off'); fig.patch.set_facecolor('white')

    ax.text(4.5, 5.25, 'Stage 2: Adaptive Noise Reduction', fontsize=13,
            fontweight='bold', ha='center', color=NAVY)
    ax.text(4.5, 4.95, 'Dual-model VAD filters noise before speech ever reaches the recognizer',
            fontsize=9, ha='center', color='#666', style='italic')

    # Raw Audio input
    rounded_box(ax, 0.3, 3.6, 2.0, 1.0,
                ['Raw Audio Input', 'Noisy military radio', '16kHz mono PCM'], LIGHT, NAVY)
    arrow_right(ax, 2.3, 4.1, 2.8)

    # Decision: Silero VAD available?
    diamond_x, diamond_y = 3.8, 4.1
    diamond = plt.Polygon([
        [diamond_x, diamond_y + 0.5],
        [diamond_x + 0.7, diamond_y],
        [diamond_x, diamond_y - 0.5],
        [diamond_x - 0.7, diamond_y],
    ], facecolor='#FFF3E0', edgecolor=ORANGE, lw=1.5)
    ax.add_patch(diamond)
    ax.text(diamond_x, diamond_y, 'Silero\navailable?', fontsize=7.5, ha='center', va='center', fontweight='bold')

    # YES → Silero Neural VAD
    ax.text(5.0, 4.5, 'Yes', fontsize=8, color=GREEN, fontweight='bold')
    arrow_right(ax, 4.5, 4.3, 5.2, color=GREEN)

    rounded_box(ax, 5.2, 3.5, 3.4, 1.2,
                ['Silero Neural VAD (Primary)', 'Deep neural network speech classifier',
                 'P(speech) computed per 32ms frame',
                 'Threshold: 0.5 (tuned for radio SNR)'],
                '#E8F8E8', GREEN, title_size=9.5, body_size=7.5)

    # NO → Energy Fallback
    ax.text(3.3, 3.2, 'No', fontsize=8, color=RED, fontweight='bold')
    arrow_down(ax, 3.8, 3.6, 2.7, color=RED)

    rounded_box(ax, 2.4, 1.5, 3.0, 1.2,
                ['Custom Energy-Based VAD (Fallback)', 'RMS energy per frame vs. adaptive floor',
                 'Ensures system never silently fails',
                 'Graceful degradation guarantee'],
                '#E8F8E8', GREEN, title_size=9.5, body_size=7.5)

    # Both paths merge → Constraints
    arrow_right(ax, 8.6, 4.1, 8.6)  # dummy
    arrow_down(ax, 6.9, 3.5, 2.7)
    arrow_right(ax, 5.4, 2.1, 5.8)

    rounded_box(ax, 5.8, 1.3, 3.0, 1.4,
                ['Custom Safety Constraints', 'Min speech duration: 250ms',
                 'Max speech duration: 30s',
                 'Pre-speech buffer: 200ms (ring buffer)',
                 'Silence gap: 600ms'],
                '#E8F8E8', GREEN, title_size=9.5, body_size=7.5)

    # Output
    arrow_down(ax, 5.0, 1.5, 0.8)
    arrow_down(ax, 7.3, 1.3, 0.8)

    rounded_box(ax, 2.5, 0.1, 5.0, 0.65,
                ['Output: Clean speech frames only (noise discarded)', ''],
                '#FFF3E0', ORANGE, title_size=9.5)

    plt.tight_layout()
    p = os.path.join(OUT_DIR, '_d2.png')
    plt.savefig(p, dpi=220, bbox_inches='tight', facecolor='white'); plt.close()
    return p


# ═══════════════════════════════════════════════════════════
# DIAGRAM 3: Speech/Non-Speech Segmentation
# ═══════════════════════════════════════════════════════════
def diagram_speech_segmentation():
    fig, ax = plt.subplots(figsize=(9, 5))
    ax.set_xlim(0, 9); ax.set_ylim(0, 5)
    ax.axis('off'); fig.patch.set_facecolor('white')

    ax.text(4.5, 4.75, 'Stage 3: Speech / Non-Speech Separation', fontsize=13,
            fontweight='bold', ha='center', color=NAVY)
    ax.text(4.5, 4.45, 'Stateful detection that produces clean, complete utterance segments',
            fontsize=9, ha='center', color='#666', style='italic')

    # State machine
    # NOT_SPEAKING state
    rounded_box(ax, 0.3, 2.8, 2.2, 1.2,
                ['NOT_SPEAKING', 'Waiting for speech', 'Monitoring each frame', 'Buffer stores last 200ms'],
                '#E0E7FF', '#4455AA', title_size=10, body_size=7.5)

    # Arrow NOT_SPEAKING → SPEAKING
    ax.annotate('', xy=(3.0, 3.7), xytext=(2.5, 3.7),
                arrowprops=dict(arrowstyle='->', color=GREEN, lw=2))
    ax.text(2.75, 3.95, 'P(speech) > 0.5', fontsize=7, ha='center', color=GREEN, fontweight='bold')

    # SPEAKING state
    rounded_box(ax, 3.0, 2.8, 2.4, 1.2,
                ['SPEAKING', 'Accumulating audio', 'Counting silence frames', 'Enforcing max 30s cap'],
                '#E8F8E8', GREEN, title_size=10, body_size=7.5)

    # Arrow SPEAKING → back to NOT_SPEAKING (curved below)
    ax.annotate('', xy=(2.0, 2.8), xytext=(4.2, 2.8),
                arrowprops=dict(arrowstyle='->', color=RED, lw=1.5,
                               connectionstyle='arc3,rad=0.4'))
    ax.text(3.1, 2.2, 'Silence > 600ms', fontsize=7, ha='center', color=RED, fontweight='bold')

    # EMIT segment
    arrow_right(ax, 5.4, 3.4, 5.9, color=NAVY, lw=2)

    rounded_box(ax, 5.9, 2.6, 2.8, 1.6,
                ['Emit Speech Segment', '', 'Audio data + timestamps',
                 'Duration (validated)',
                 'Confidence score',
                 'Saved as WAV for audit'],
                '#FFF3E0', ORANGE, title_size=10, body_size=7.5)

    # Key features at bottom
    features = [
        ('Pre-speech padding', 'Ring buffer captures 200ms before VAD trigger, so word onsets are never clipped'),
        ('Min duration filter', 'Segments shorter than 250ms are discarded (rejects clicks, key noise, static bursts)'),
        ('Max duration cap', 'Segments exceeding 30s are force-split to prevent memory overflow'),
        ('Per-segment WAV', 'Every segment is saved as an individual WAV file for audit trail and replay'),
    ]
    y = 1.7
    for title, desc in features:
        ax.text(0.5, y, title + ':', fontsize=7.5, fontweight='bold', color=GREEN, va='center')
        ax.text(2.8, y, desc, fontsize=7.5, color=DARK, va='center')
        y -= 0.35

    plt.tight_layout()
    p = os.path.join(OUT_DIR, '_d3.png')
    plt.savefig(p, dpi=220, bbox_inches='tight', facecolor='white'); plt.close()
    return p


# ═══════════════════════════════════════════════════════════
# DIAGRAM 4: Word Segmentation & Post-Processing
# ═══════════════════════════════════════════════════════════
def diagram_word_segmentation():
    fig, ax = plt.subplots(figsize=(9, 5.8))
    ax.set_xlim(0, 9); ax.set_ylim(0, 5.8)
    ax.axis('off'); fig.patch.set_facecolor('white')

    ax.text(4.5, 5.55, 'Stage 4: Word Segmentation & Domain Post-Processing', fontsize=13,
            fontweight='bold', ha='center', color=NAVY)
    ax.text(4.5, 5.25, '6-stage custom pipeline that turns raw Whisper output into accurate military transcription',
            fontsize=9, ha='center', color='#666', style='italic')

    # 6 pipeline stages as vertical steps
    stages = [
        ('1', 'Whisper STT + Domain Prompt', 'Faster-Whisper large-v3-turbo with 46 injected\nmilitary terms biasing the decoder', '#E8F8E8', GREEN),
        ('2', 'Filler Word Removal', 'Strips hesitations: "uh", "um", "hmm", "er",\n"you know", "i mean", "like"', '#E8F8E8', GREEN),
        ('3', 'Domain Correction Dictionary', '50+ rules mapping Whisper errors to correct terms\ne.g. "fuel presser" -> "fuel pressure"', '#E8F8E8', GREEN),
        ('4', 'Abbreviation Expansion', '20 military acronyms: INS, FTS, LOX, UDMH,\nIMU, OBC, GCS, RSO, MET, MECO...', '#E8F8E8', GREEN),
        ('5', 'Number Normalization', 'Converts spoken numbers to digits:\n"twenty three psi" -> "23 psi"', '#E8F8E8', GREEN),
        ('6', 'STT Artifact Deduplication', 'Removes consecutive repeated words\n(common Whisper failure mode)', '#E8F8E8', GREEN),
    ]

    x_left = 0.3
    box_w = 5.0
    box_h = 0.62
    gap = 0.08
    start_y = 4.65

    for i, (num, title, desc, bg, edge) in enumerate(stages):
        y = start_y - i * (box_h + gap)
        box = FancyBboxPatch((x_left, y), box_w, box_h, boxstyle="round,pad=0.08",
                              facecolor=bg, edgecolor=edge, linewidth=1.5)
        ax.add_patch(box)
        # Number circle
        circle = plt.Circle((x_left + 0.3, y + box_h/2), 0.18, facecolor=edge, edgecolor='white', lw=1)
        ax.add_patch(circle)
        ax.text(x_left + 0.3, y + box_h/2, num, fontsize=9, fontweight='bold',
                ha='center', va='center', color='white')
        # Title
        ax.text(x_left + 0.65, y + box_h/2 + 0.05, title, fontsize=8.5, fontweight='bold',
                ha='left', va='center', color='#1B6B1B')
        # Description (smaller, to the right)
        ax.text(x_left + 0.65, y + box_h/2 - 0.17, desc, fontsize=6.5,
                ha='left', va='center', color=DARK)
        # Arrow down
        if i < len(stages) - 1:
            arrow_down(ax, x_left + 2.5, y, y - gap, color=NAVY, lw=1.2)

    # Before / After example on the right
    ex_x = 5.7
    # BEFORE
    box = FancyBboxPatch((ex_x, 3.4), 3.0, 1.5, boxstyle="round,pad=0.1",
                          facecolor='#FFF0F0', edgecolor=RED, linewidth=1.5)
    ax.add_patch(box)
    ax.text(ex_x + 1.5, 4.7, 'BEFORE', fontsize=9, fontweight='bold', ha='center', color=RED)
    ax.text(ex_x + 1.5, 4.45, '(Raw Whisper output)', fontsize=7, ha='center', color='#999')
    ax.text(ex_x + 0.15, 4.1, '"uh fuel presser uh\n tell a metric um nominal\n twenty three psi\n confirmed confirmed"',
            fontsize=7, ha='left', va='center', color=DARK, family='monospace')

    # Arrow
    arrow_down(ax, ex_x + 1.5, 3.4, 2.9, color=GREEN, lw=2)

    # AFTER
    box = FancyBboxPatch((ex_x, 1.6), 3.0, 1.3, boxstyle="round,pad=0.1",
                          facecolor='#E8F8E8', edgecolor=GREEN, linewidth=1.5)
    ax.add_patch(box)
    ax.text(ex_x + 1.5, 2.75, 'AFTER', fontsize=9, fontweight='bold', ha='center', color=GREEN)
    ax.text(ex_x + 1.5, 2.5, '(Our post-processing)', fontsize=7, ha='center', color='#999')
    ax.text(ex_x + 0.15, 2.05, '"fuel pressure\n telemetry nominal\n 23 psi confirmed"',
            fontsize=7, ha='left', va='center', color=DARK, family='monospace')

    # Bottom summary
    ax.text(4.5, 0.6, 'Result: Clean, domain-accurate transcription with word-level timestamps',
            fontsize=9, ha='center', va='center', color=NAVY, fontweight='bold')
    ax.text(4.5, 0.25, '5 custom stages | 50+ correction rules | 20 abbreviation expansions | Number conversion',
            fontsize=8, ha='center', va='center', color='#666', style='italic')

    plt.tight_layout()
    p = os.path.join(OUT_DIR, '_d4.png')
    plt.savefig(p, dpi=220, bbox_inches='tight', facecolor='white'); plt.close()
    return p


# ═══════════════════════════════════════════════════════════
# BUILD WORD DOCUMENT
# ═══════════════════════════════════════════════════════════
def build_document():
    print("Generating diagrams...")
    d1 = diagram_overview()
    d2 = diagram_noise_reduction()
    d3 = diagram_speech_segmentation()
    d4 = diagram_word_segmentation()
    print("Diagrams done.")

    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21); sec.page_height = Cm(29.7)
    sec.top_margin = Cm(1.5); sec.bottom_margin = Cm(1.5)
    sec.left_margin = Cm(1.8); sec.right_margin = Cm(1.8)

    # ════════════════ HEADER ════════════════
    para(doc, 'DEFENCE RESEARCH AND DEVELOPMENT ORGANISATION (DRDO)',
         size=11, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, 'Pre-Launch Checklist Verification System (PLCVS)',
         size=14, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    para(doc, 'Technical Report: Audio Processing Pipeline',
         size=12, bold=True, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, 'Noise Reduction  |  Speech/Non-Speech Separation  |  Word Segmentation',
         size=10, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)

    # Separator
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run('_' * 90)
    r.font.size = Pt(5); r.font.color.rgb = RGBColor.from_string('CCCCCC')

    # ════════════════ 1. INTRODUCTION ════════════════
    para(doc, '1.  What PLCVS Does and Why We Built Custom Processing',
         size=11, bold=True, color=NAVY, after=4)

    para(doc,
        'PLCVS listens to military radio communication during missile pre-launch checklist '
        'verification and automatically tracks which checklist items have been confirmed or '
        'failed. The entire system runs offline on edge hardware -- no internet connection '
        'is available or allowed during live operations.',
        size=9.5, after=4)

    para(doc,
        'The challenge is that military radio audio is very different from the clean speech '
        'that commercial speech recognition systems are designed for. We deal with background '
        'noise from launch-pad environments, domain-specific terminology that standard models '
        'consistently get wrong (words like "telemetry", "umbilical", "INS alignment"), and '
        'strict real-time requirements -- the system must process and display results within '
        '5 seconds of a person speaking.',
        size=9.5, after=4)

    para(doc,
        'These constraints meant we could not simply plug in an off-the-shelf speech recognizer '
        'and expect useful results. We built a multi-stage processing pipeline with custom '
        'noise reduction, intelligent speech segmentation, and extensive domain-specific '
        'post-processing. The diagram below shows the four main stages.',
        size=9.5, after=6)

    # ── Figure 1 ──
    para(doc, 'Figure 1: End-to-End Pipeline Overview',
         size=9, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    add_image_centered(doc, d1)
    para(doc, '', size=4, after=6)

    # ════════════════ 2. NOISE REDUCTION ════════════════
    para(doc, '2.  Noise Reduction: How We Filter Out Unwanted Sound',
         size=11, bold=True, color=NAVY, after=4)

    para(doc,
        'Before any speech recognition happens, we need to separate human speech from '
        'everything else -- engine noise, wind, radio static, mechanical clicks from push-to-talk '
        'buttons, and periods of silence. Our approach uses Voice Activity Detection (VAD) '
        'to analyze audio frame-by-frame and decide: "Is someone speaking right now, or is '
        'this just noise?"',
        size=9.5, after=4)

    para(doc, 'What we built:', size=10, bold=True, after=3)

    bullet(doc,
        'Primary detector: We use the Silero neural network VAD model, which processes audio '
        'in 32-millisecond frames and outputs a probability score for each frame. We tuned '
        'the decision threshold to 0.5, specifically calibrated for the signal-to-noise ratio '
        'typical of military radio channels.', size=9)

    bullet(doc,
        'Custom energy-based fallback: If the neural model is unavailable or encounters '
        'unusual audio it was not trained on (radio codec artifacts, for example), the system '
        'automatically falls back to our custom energy-based VAD. This computes the RMS energy '
        'of each frame against an adaptive noise floor. The key point: the system never '
        'silently fails -- it always has a working detection path.', size=9)

    bullet(doc,
        'Pre-speech ring buffer: We maintain a rolling 200ms audio buffer at all times. '
        'When VAD detects speech onset, we prepend this buffer to the speech segment. '
        'Without this, the first syllable of each utterance would be clipped -- a common '
        'problem in naive VAD implementations.', size=9)

    bullet(doc,
        'Duration constraints: Minimum speech duration of 250ms filters out transient noise '
        '(a click or cough), while maximum duration of 30s prevents buffer overflow from '
        'continuous radio keying. The silence gap threshold of 600ms determines when an '
        'utterance has ended.', size=9, after=6)

    # ── Figure 2 ──
    para(doc, 'Figure 2: Noise Reduction -- Dual-Model VAD Architecture',
         size=9, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    add_image_centered(doc, d2)
    para(doc, '', size=4, after=6)

    # ════════════════ 3. SPEECH / NON-SPEECH ════════════════
    para(doc, '3.  Speech / Non-Speech Separation: Producing Clean Segments',
         size=11, bold=True, color=NAVY, after=4)

    para(doc,
        'Once we know which frames contain speech, the next challenge is grouping those frames '
        'into complete, meaningful utterance segments. A person might say "fuel pressure nominal" '
        'as a single phrase, but the VAD sees it as a sequence of individual 32ms frames. Our '
        'stateful segment detector assembles these frames into coherent speech regions with '
        'precise start and end timestamps.',
        size=9.5, after=4)

    para(doc, 'What we built:', size=10, bold=True, after=3)

    bullet(doc,
        'Finite state machine: The detector maintains two states -- NOT_SPEAKING and SPEAKING. '
        'It transitions to SPEAKING when the VAD probability exceeds 0.5, and back to '
        'NOT_SPEAKING only after sustained silence of at least 600ms. This hysteresis prevents '
        'brief pauses within a sentence (like between words) from prematurely splitting the segment.', size=9)

    bullet(doc,
        'Speech region accumulation: During the SPEAKING state, audio frames are continuously '
        'accumulated into a buffer. When the state transitions back to NOT_SPEAKING, the '
        'accumulated audio is packaged into a SpeechRegion object with its start time, end time, '
        'duration, audio data, and a confidence score derived from the average VAD probability.', size=9)

    bullet(doc,
        'Segment quality validation: Before emitting a segment, we validate that it meets our '
        'minimum duration requirement (250ms). Segments that are too short are discarded as noise. '
        'Segments that exceed the maximum duration (30s) are force-emitted to prevent unbounded '
        'memory growth.', size=9)

    bullet(doc,
        'Audit trail: Each segment is automatically saved as an individual WAV file with a '
        'timestamped filename. This provides a complete, reviewable record of every speech '
        'segment the system processed during a verification session.', size=9, after=6)

    # ── Figure 3 ──
    para(doc, 'Figure 3: Speech Segmentation State Machine',
         size=9, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    add_image_centered(doc, d3)
    para(doc, '', size=4, after=6)

    # ════════════════ 4. WORD SEGMENTATION ════════════════
    para(doc, '4.  Word Segmentation & Domain-Specific Post-Processing',
         size=11, bold=True, color=NAVY, after=4)

    para(doc,
        'The speech segments are transcribed using Faster-Whisper (large-v3-turbo model), which '
        'provides word-level timestamps through forced alignment. However, raw Whisper output '
        'on military radio audio contains significant errors -- domain terms are misrecognized, '
        'filler words clutter the text, and numbers are spelled out rather than digitized. '
        'We built a custom 6-stage post-processing pipeline to address each of these problems.',
        size=9.5, after=4)

    para(doc, 'What we built:', size=10, bold=True, after=3)

    bullet(doc,
        'Domain vocabulary injection: We feed a curated prompt of 46 military terms (fuel pressure, '
        'telemetry, INS alignment, flight termination system, etc.) into Whisper\'s decoder as an '
        'initial prompt. This biases the language model toward domain-specific vocabulary and '
        'significantly reduces misrecognition of critical terminology. The vocabulary is stored '
        'in an external YAML file, making it configurable per mission type.', size=9)

    bullet(doc,
        'Domain correction dictionary: We maintain a dictionary of 50+ common Whisper '
        'misrecognitions specific to our domain. For example, Whisper often transcribes '
        '"telemetry" as "tell a metric" or "tell em a tree", "umbilical" as "umber lick", '
        'and "INS" as "eye and s". Our correction rules are applied in length-sorted order '
        '(longest match first) to handle overlapping patterns correctly.', size=9)

    bullet(doc,
        'Military abbreviation expansion: 20 defence acronyms (INS, FTS, LOX, UDMH, IMU, OBC, '
        'GCS, RSO, MET, MECO, etc.) are recognized and can be expanded to their full forms, '
        'enabling downstream NLP processing to handle both spoken and abbreviated forms.', size=9)

    bullet(doc,
        'Spoken number normalization: Converts verbal number expressions to digits -- '
        '"twenty three" becomes "23", "one hundred" becomes "100". This is essential for '
        'pressure readings, voltage levels, and countdown values that appear throughout '
        'checklist verification.', size=9)

    bullet(doc,
        'Artifact removal: Consecutive duplicate words (a common Whisper failure where it '
        'outputs "confirmed confirmed") are deduplicated. Vocal fillers ("uh", "um", "hmm", '
        '"you know") are stripped. The result is clean, publication-ready transcription text.', size=9, after=6)

    # ── Figure 4 ──
    para(doc, 'Figure 4: Word Segmentation & Post-Processing Pipeline',
         size=9, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    add_image_centered(doc, d4)
    para(doc, '', size=4, after=6)

    # ════════════════ 5. CONTRIBUTION SUMMARY ════════════════
    para(doc, '5.  Summary of Original Contributions',
         size=11, bold=True, color=NAVY, after=4)

    para(doc,
        'The table below summarizes what we built versus what comes from existing libraries. '
        'The core insight is that while we leverage established models (Whisper for STT, Silero '
        'for VAD), the integration logic, domain adaptation, fault tolerance, and post-processing '
        'are entirely our own design and implementation.',
        size=9.5, after=6)

    # Table
    rows_data = [
        ('Dual-Model VAD Architecture',
         'Neural VAD with custom energy-based fallback; adaptive threshold tuning for military radio SNR',
         'Zero silent failures; works on any hardware'),
        ('Pre-Speech Ring Buffer',
         '200ms rolling buffer that preserves word onsets before VAD trigger fires',
         'No syllable clipping at segment start'),
        ('Stateful Speech Detector',
         'FSM with hysteresis, min/max duration enforcement, silence gap detection',
         'Clean, complete utterance segments'),
        ('Domain Vocabulary Injection',
         '46-term military prompt fed into Whisper decoder; externalized YAML config',
         'Significant accuracy gain on domain terms'),
        ('Correction Dictionary',
         '50+ rules mapping Whisper misrecognitions to correct military terms',
         'Handles terms no retraining could fix'),
        ('Post-Processing Pipeline',
         '6-stage pipeline: filler removal, corrections, abbreviation expansion, number normalization, deduplication',
         'Clean text from noisy radio audio'),
    ]

    table = doc.add_table(rows=len(rows_data)+1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, h in enumerate(['Component', 'What We Built', 'Impact']):
        c = table.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(9); r.bold = True
                r.font.color.rgb = RGBColor(255,255,255); r.font.name = 'Calibri'
        shade_cell(c, NAVY)

    for ri, (comp, contrib, impact) in enumerate(rows_data, 1):
        table.rows[ri].cells[0].text = comp
        table.rows[ri].cells[1].text = contrib
        table.rows[ri].cells[2].text = impact
        for ci in range(3):
            for p in table.rows[ri].cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8); r.font.name = 'Calibri'
            if ri % 2 == 0:
                shade_cell(table.rows[ri].cells[ci], 'F5F5F5')

    for row in table.rows:
        row.cells[0].width = Cm(3.2)
        row.cells[1].width = Cm(9.0)
        row.cells[2].width = Cm(5.0)

    para(doc, '', size=6, after=8)

    para(doc,
        'PLCVS is not a wrapper around existing libraries. It is a purpose-built defence system '
        'with custom signal processing, domain adaptation layers, and a robust post-processing '
        'pipeline -- all designed to solve the specific challenges of real-time military radio '
        'speech processing in offline, mission-critical environments.',
        size=10, bold=True, color=NAVY, after=8)

    # Footer separator
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run('_' * 90)
    r.font.size = Pt(5); r.font.color.rgb = RGBColor.from_string('CCCCCC')

    para(doc, 'DRDO - Pre-Launch Checklist Verification System  |  Technical Report  |  April 2026',
         size=8, color='999999', align=WD_ALIGN_PARAGRAPH.CENTER, after=0)

    # Save
    out = os.path.join(OUT_DIR, 'PLCVS_Technical_Report_Audio_Pipeline.docx')
    doc.save(out)
    print(f"\nSaved: {out}")

    # Cleanup
    for img in [d1, d2, d3, d4]:
        os.remove(img)
    print("Done.")
    return out


if __name__ == '__main__':
    build_document()
